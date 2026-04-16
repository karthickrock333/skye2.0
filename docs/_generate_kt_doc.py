"""Generate KT_ARCHITECTURE.docx from structured content."""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ───────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0xCB, 0x20, 0x26)  # Hitachi red
    hs.font.name = 'Calibri'

BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x33, 0x33, 0x33)
RED = RGBColor(0xCB, 0x20, 0x26)
GRAY = RGBColor(0x66, 0x66, 0x66)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK


def add_table(headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing


def add_bullet(text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)


def add_para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


# ═════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('HD SKYE Agent', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RED
    run.font.size = Pt(36)

subtitle = doc.add_heading('Architecture & Knowledge Transfer Document', level=0)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = BLUE
    run.font.size = Pt(20)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('HD SKYE Agentic HR RAG Agent v2.0\n').font.size = Pt(14)
meta.add_run('Prepared for KT Session — April 2026\n').font.size = Pt(12)
meta.add_run('AI-Powered HR Policy Assistant for Hitachi Digital').font.size = Pt(12)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. System Overview",
    "2. High-Level Architecture",
    "3. Agent Variants (Main, PCard, Bulk Expense, Payroll)",
    "4. Agentic Pipeline — Main Orchestrator",
    "5. P-Card Pipeline — Dedicated Orchestrator",
    "6. Individual Agent Breakdown",
    "7. Tools Layer",
    "8. Data Architecture & Index Groups",
    "9. ServiceNow Integration",
    "10. Document Ingestion Pipeline",
    "11. Caching Strategy",
    "12. Access Control & Role-Based Architecture",
    "13. Infrastructure & Deployment",
    "14. API Endpoints",
    "15. Key Configuration Reference",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════
# SECTION 1: System Overview
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('1. System Overview', level=1)

add_para(
    'HD SKYE is an agentic Retrieval-Augmented Generation (RAG) system that serves as an '
    'HR policy assistant for Hitachi Digital employees. It answers HR-related questions '
    '(leave policies, payroll, travel, benefits, etc.) by:'
)
add_bullet('Understanding the user\'s question (language, intent, region, role)')
add_bullet('Retrieving relevant policy documents from vectorized knowledge bases')
add_bullet('Generating contextual, role-aware, region-specific answers using Google Gemini LLM')
add_bullet('Translating responses to the user\'s language (supports 15+ languages)')

doc.add_heading('Key Capabilities', level=2)
add_bullet('Multi-language support — detects language, translates queries to English, translates answers back')
add_bullet('Role-aware responses — different answers for Employee, Manager, VP, Executive, HR/Finance')
add_bullet('Region-specific policies — serves country-specific HR policies (India, Japan, US, Germany, etc.)')
add_bullet('Multi-variant agents — specialized agents for P-Card, Bulk Expense, Payroll alongside the main HR agent')
add_bullet('Conversation memory — multi-turn conversations with context preservation')
add_bullet('Employee lookup — Managers/VPs can ask about their reportees\' policies')
add_bullet('ServiceNow KB integration — ingests and serves ServiceNow knowledge base articles')

doc.add_heading('Tech Stack', level=2)
add_table(
    ['Component', 'Technology'],
    [
        ['Backend', 'Python / FastAPI'],
        ['Frontend', 'React (Vite + Bun)'],
        ['LLM', 'Google Gemini 2.0 Flash (Vertex AI)'],
        ['Vector DB', 'Google Vertex AI Matching Engine'],
        ['Document Store', 'Google Cloud Firestore (multi-database)'],
        ['Embeddings', 'Vertex AI text-embedding-004'],
        ['Cache', 'Redis (conversation history + semantic cache)'],
        ['User Data', 'BigQuery (hd-onedata-prod)'],
        ['Object Storage', 'Google Cloud Storage (GCS)'],
        ['OCR', 'Google Document AI'],
        ['Translation', 'Google Cloud Translation API + Gemini LLM'],
        ['Agent Framework', 'Google ADK (Agent Development Kit)'],
        ['Deployment', 'Cloud Run (Docker, 2 CPU / 2 GB)'],
    ],
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 2: High-Level Architecture
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('2. High-Level Architecture', level=1)

add_para('The system is a monolithic Python application serving both the React frontend and all API endpoints from a single Cloud Run service.')

add_code_block(
    '┌───────────────────────────────────────────────────────────────────────┐\n'
    '│                         FRONTEND (React SPA)                         │\n'
    '│               Served from same origin via FastAPI static              │\n'
    '└────────────────────────────────┬──────────────────────────────────────┘\n'
    '                                 │ HTTP POST\n'
    '                                 ▼\n'
    '┌───────────────────────────────────────────────────────────────────────┐\n'
    '│                       FastAPI Backend (main.py)                       │\n'
    '│                                                                       │\n'
    '│   /chat ──────────────────► Main Orchestrator (process_query)         │\n'
    '│   /pcard/chat ────────────► P-Card Orchestrator (process_pcard_query) │\n'
    '│   /bulk-expense/chat ─────► Main Orchestrator (variant=bulk_expense)  │\n'
    '│   /payroll/chat ──────────► Main Orchestrator (variant=payroll)       │\n'
    '│   /feedback ──────────────► Feedback Agent                            │\n'
    '│   /documents/{file} ──────► GCS Signed URL / KB Renderer              │\n'
    '└──────────────────────────────────────────────────────────────────────┘\n'
    '                                 │\n'
    '           ┌─────────────────────┼─────────────────────┐\n'
    '           ▼                     ▼                     ▼\n'
    '┌──────────────────┐  ┌──────────────────┐  ┌──────────────────────────┐\n'
    '│  ORCHESTRATOR    │  │  AGENT MODULES   │  │  EXTERNAL SERVICES       │\n'
    '│  (Pipeline Ctrl) │  │  (15 Agents)     │  │                          │\n'
    '│  • Phase control │  │  • Query Under.  │  │  • Vertex AI (LLM)       │\n'
    '│  • Thread pools  │  │  • Guardrails    │  │  • Matching Engine       │\n'
    '│  • Cache checks  │  │  • Access Ctrl   │  │  • Firestore (5 DBs)     │\n'
    '│  • Role routing  │  │  • Retrieval     │  │  • BigQuery (user data)  │\n'
    '│  • Fallback      │  │  • Reranking     │  │  • Redis (cache)         │\n'
    '│                  │  │  • Generation    │  │  • GCS (documents)       │\n'
    '│                  │  │  • Translation   │  │  • Cloud Translation     │\n'
    '│                  │  │  • Post Valid.   │  │  • Document AI (OCR)     │\n'
    '└──────────────────┘  │  • Caching       │  └──────────────────────────┘\n'
    '                      │  • Observability │\n'
    '                      └──────────────────┘'
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 3: Agent Variants
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Agent Variants (Main, PCard, Bulk Expense, Payroll)', level=1)

add_para(
    'SKYE supports 4 named agent variants, each configured with a different combination '
    'of vector index groups. This allows specialized agents to focus on specific policy '
    'domains while sharing the same codebase.'
)

doc.add_heading('Variant Registry', level=2)
add_table(
    ['Variant', 'Display Name', 'Index Groups Searched', 'Priority Groups', 'Dedicated Route'],
    [
        ['main', 'Skye HR Agent', 'servicenow_kb, main, apac_payroll, pcard, bulk_exp', 'None (all equal)', 'POST /chat'],
        ['pcard', 'Skye P-Card Agent', 'servicenow_kb, pcard', 'pcard (boosted)', 'POST /pcard/chat'],
        ['bulk_expense', 'Skye Bulk Expense Agent', 'servicenow_kb, bulk_exp', 'bulk_exp (boosted)', 'POST /bulk-expense/chat'],
        ['payroll', 'Skye Payroll Agent', 'apac_payroll', 'apac_payroll (boosted)', 'POST /payroll/chat'],
    ],
)

doc.add_heading('How Variants Work', level=2)
add_para('When a request arrives with a specific variant, the system:', bold=True)
add_bullet('Looks up AgentVariantConfig with the list of index group names', bold_prefix='1. Config Lookup: ')
add_bullet('Resolves to actual IndexGroupConfig objects (endpoint IDs, Firestore DBs, collections)', bold_prefix='2. Index Resolution: ')
add_bullet('Retrieval agent searches ONLY those resolved indexes', bold_prefix='3. Scoped Search: ')
add_bullet('Returns the Firestore collection names for priority groups', bold_prefix='4. Priority Collections: ')
add_bullet('Reranking applies INDEX_PRIORITY_BOOST (+0.20 score) to results from priority collections', bold_prefix='5. Reranking Boost: ')

doc.add_heading('Variant-Specific Behavior', level=2)
add_bullet('Searches ALL indexes, no priority boost. Full pipeline with access control, role rewriting, region filtering.', bold_prefix='Main: ')
add_bullet('Has its own dedicated orchestrator (pcard_orchestrator.py) with strict gold-source rules, no access control, custom LLM prompt, fallback to CorporateCard@hitachidigital.com.', bold_prefix='P-Card: ')
add_bullet('Uses main orchestrator with variant="bulk_expense", prioritizes bulk expense index.', bold_prefix='Bulk Expense: ')
add_bullet('Uses main orchestrator with variant="payroll", searches only APAC payroll index.', bold_prefix='Payroll: ')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 4: Main Orchestrator Pipeline
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Agentic Pipeline — Main Orchestrator', level=1)

add_para(
    'The main orchestrator (agents/orchestrator.py → process_query()) uses aggressive '
    'parallelization with ThreadPoolExecutor to overlap slow operations. The pipeline '
    'is divided into 5 phases.'
)

doc.add_heading('Phase 1: Launch ALL Independent Work at t=0', level=2)
add_para('Four threads launch simultaneously:')
add_bullet('Fetch conversation history (instant, ~10ms)', bold_prefix='Thread 1: ')
add_bullet('Home-location access check (runs in background)', bold_prefix='Thread 2: ')
add_bullet('Speculative vector retrieval (runs in background, ~2-3s)', bold_prefix='Thread 3: ')
add_bullet('Generate query embedding for semantic cache (~0.5s)', bold_prefix='Thread 4: ')
add_para(
    'The main thread waits ONLY for history, then immediately runs Query Understanding '
    '(~1-2s) which does: language detection, translation, intent classification, '
    'region extraction, employee name detection, abbreviation expansion, and follow-up rewriting.',
    bold=False,
)
add_para('Early exits at this phase:', bold=True)
add_bullet('Translation-only request → retranslate last answer without re-running pipeline')

doc.add_heading('Guardrails (Inline, Fast)', level=2)
add_bullet('Greeting/thank-you detection → canned response, skip pipeline')
add_bullet('Hitachi Vantara (HV) blocking → disclaimer')
add_bullet('P-Card permission gating → VP/Exec/SuperAdmin/HR-Finance only')
add_para('If blocked → return early, cancel background threads.')

doc.add_heading('Cache Check', level=2)
add_para('Two-level cache check before proceeding to retrieval:')
add_bullet('Exact match: answer_cache:{variant}:{query_en}:{region} — pure Redis, <10ms', bold_prefix='Level 1: ')
add_bullet('Semantic similarity: cosine sim ≥ 0.85 on cached embeddings, role-aware bucketing (employee/manager/vp/executive), intent validation (reject if cached intent ≠ current intent)', bold_prefix='Level 2: ')
add_bullet('Denial/error/fallback cached answers are never served (stale denial protection)')

doc.add_heading('Phase 2: Access Control', level=2)
add_bullet('Employee lookup (if mentioned): find in reportees, validate access')
add_bullet('Role-based region permission matrix check')
add_bullet('If denied → soft deny: fall back to user\'s home region (not hard block)')
add_bullet('Role-based query rewriting — augment search with role-specific terms')
add_bullet('Region augmentation: append "in {target_region}" to search query')

doc.add_heading('Phase 3: Collect Retrieval Results', level=2)
add_bullet('Collect speculative retrieval (launched at t=0, should be ready by now)')
add_bullet('If holiday/translated/region-specific: launch REFINED retrieval with English search_query + region context → merge results (dedup by chunk ID)')
add_bullet('If P-Card authorized → branch to P-Card sub-pipeline inside main orchestrator')

doc.add_heading('Reranking + Generation (Sequential)', level=2)
add_para('Reranking stages:', bold=True)
add_bullet('HV source filtering')
add_bullet('Index priority boost (variant-specific, +0.20)')
add_bullet('Category boost (P-Card, Bulk Expense)')
add_bullet('Holiday/PCard result prioritization')
add_bullet('Region filtering: multi-method country detection (Firestore metadata → filename → text scan)')
add_bullet('OPCO entity classification (HDS, GlobalLogic, HD, Global)')
add_para('Generation (Gemini 2.0 Flash):', bold=True)
add_bullet('12-rule master prompt with role context, region, OPCO labels')
add_bullet('Concise vs Detailed mode based on intent')
add_bullet('If "no info" fallback → retry with broader search (top_k=50, global scope)')

doc.add_heading('Phase 5: Parallel Post-Processing', level=2)
add_bullet('Translation — translate English answer to user\'s language', bold_prefix='Thread 1: ')
add_bullet('Post-Validation — attribute sources, detect no-info fallbacks', bold_prefix='Thread 2: ')
add_bullet('Follow-up Suggestions — generate 3 suggested questions', bold_prefix='Thread 3: ')
add_para('Then: build source links, cache answer (exact + semantic, 3h TTL), cache session context (24h), save conversation turn to Redis.')

doc.add_heading('Parallelization Speedup', level=2)
add_table(
    ['Sequential (old)', 'Parallel (current)', 'Savings'],
    [
        ['Understanding (2s) → Retrieval (3s)', 'Understanding ‖ Speculative Retrieval', '~3s saved'],
        ['Translation → Validation → Suggestions', 'Translation ‖ Validation ‖ Suggestions', '~2-3s saved'],
        ['History → Home-loc → Embedding', 'All at t=0 in threads', '~1s saved'],
    ],
)
add_para(
    'Speculative Retrieval: Retrieval starts at t=0 using the raw question, before '
    'understanding completes. If the question is non-English or region-specific, a '
    'refined retrieval runs after understanding with the corrected English query + region context, '
    'and results are merged.',
    italic=True,
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 5: P-Card Pipeline
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('5. P-Card Pipeline — Dedicated Orchestrator', level=1)

add_para('The P-Card variant has its own orchestrator (pcard_orchestrator.py) with fundamentally different logic.')

doc.add_heading('Differences from Main Pipeline', level=2)
add_table(
    ['Aspect', 'Main Pipeline', 'P-Card Pipeline'],
    [
        ['Access Control', 'Full role-based', 'None — open to all'],
        ['Region Filtering', 'Country-specific', 'None — global policy'],
        ['Role Rewriting', 'Role-aware query augmentation', 'None'],
        ['Source Filtering', 'HV filter + region filter', 'Strict gold-source only'],
        ['LLM Prompt', 'General HR policy expert', 'Procurement Card Policy Expert'],
        ['Fallback', 'AskNow + HRBP redirect', 'CorporateCard@hitachidigital.com'],
        ['Footnote Handling', 'N/A', '(*) → Gift Policy, (**) → 3P Policy'],
    ],
)

doc.add_heading('Gold Source Rule', level=2)
add_bullet('ONLY use PCard_Allowable_NonAllowable.png (the master table)', bold_prefix='Rule 1: ')
add_bullet('If the table has (*) marker → include Global Employee Gift Policy PDF', bold_prefix='Rule 2: ')
add_bullet('If the table has (**) marker → include Third Party Gifts/Travel/Entertainment Policy PDF', bold_prefix='Rule 3: ')
add_bullet('NO other sources allowed', bold_prefix='Rule 4: ')
add_bullet('If no PNG chunk found → return fallback email message', bold_prefix='Rule 5: ')

doc.add_heading('P-Card Access from Main Pipeline (/chat route)', level=2)
add_para(
    'When a user sends a P-Card query to /chat (main route), the guardrails agent detects it '
    'and checks permissions:'
)
add_bullet('If user is VP/Executive/SuperAdmin/HR-Finance → P-Card sub-pipeline runs inside the main orchestrator')
add_bullet('If user is a regular employee → denied with a role-based message')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 6: Individual Agent Breakdown
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Individual Agent Breakdown', level=1)

# 6.1
doc.add_heading('6.1 Query Understanding Agent', level=2)
add_para('File: agents/query_understanding_agent.py', italic=True, color=GRAY)
add_para('Purpose: First stage — parses raw user input into structured understanding.')
add_table(
    ['Output Field', 'Description'],
    [
        ['query_en', 'English translation of the query'],
        ['response_language_code', 'Target language for response (e.g., "ja", "ta")'],
        ['is_greeting', 'True if query is a greeting/thank-you'],
        ['is_followup', 'True if query depends on conversation history'],
        ['intent', '"concise" or "detailed"'],
        ['search_query', 'Expanded/rewritten query for vector search'],
        ['target_region', 'Country the query is about (e.g., "India", "Japan")'],
        ['mentioned_employee', 'Employee name if asking about a reportee'],
    ],
)
add_para('Key Algorithms:', bold=True)
add_bullet('Single combined LLM call: Does follow-up detection + intent classification + query rewriting + geographic extraction + employee detection + translation correction in ONE Gemini call')
add_bullet('Query expansion dictionary: 150+ entries mapping HR abbreviations ("wfh" → "work from home")')
add_bullet('Follow-up rewriting: Converts context-dependent queries to standalone form')

# 6.2
doc.add_heading('6.2 Guardrails Agent', level=2)
add_para('File: agents/guardrails_agent.py', italic=True, color=GRAY)
add_para('Purpose: Pre-flight permission gate — early exit for blocked/greeting queries.')
add_bullet('Greeting guard → canned response, skip pipeline', bold_prefix='Guard 1: ')
add_bullet('HV blocking → Hitachi Vantara queries get disclaimer', bold_prefix='Guard 2: ')
add_bullet('P-Card gating → checks VP/Exec/SuperAdmin/HR-Finance role via BQ', bold_prefix='Guard 3: ')

# 6.3
doc.add_heading('6.3 Access Control Agent', level=2)
add_para('File: agents/access_control_agent.py', italic=True, color=GRAY)
add_para('Purpose: Role-based region access matrix.')
add_table(
    ['Role', 'Can Access', 'Employee Lookup'],
    [
        ['Regular Employee', 'Global + own country', 'Denied'],
        ['Manager', 'Global + own + direct reports\' countries', '1 level deep'],
        ['VP / Executive', 'Global + own + all reports\' countries', '2 levels deep'],
        ['Super Admin', 'ALL regions (bypass)', 'No (unless also Manager/VP)'],
        ['HR / Finance', 'ALL regions', 'Standard per other role'],
        ['data_scope=global', 'ALL regions (override)', 'Standard per other role'],
    ],
)
add_para('Country Resolution: Fast static map for 99% of cases (e.g., "Hyderabad" → "India"). LLM fallback for unknown locations, cached for 30 days.')

# 6.4
doc.add_heading('6.4 Retrieval Agent', level=2)
add_para('File: agents/retrieval_agent.py', italic=True, color=GRAY)
add_para('Purpose: Vector similarity search against Matching Engine indexes.')
add_bullet('Multi-index parallel search: Groups indexes by endpoint, searches all on each endpoint in parallel')
add_bullet('Rich metadata extraction: country, region, category, document_type, servicenow_url, etc.')
add_bullet('Semantic re-ranking: Uses Vertex AI Ranking API (semantic-ranker-512@latest)')
add_bullet('Variant-aware caching: 1-hour TTL with variant fingerprints')

# 6.5
doc.add_heading('6.5 Reranking Agent', level=2)
add_para('File: agents/reranking_agent.py', italic=True, color=GRAY)
add_para('Purpose: Post-retrieval multi-stage ranking and filtering.')
add_bullet('Stage 1: HV Filtering — Remove Hitachi Vantara results')
add_bullet('Stage 1.5: Index Priority Boost — +0.20 score for variant\'s priority indexes')
add_bullet('Stage 1.6: Category Boost — Finer-grained boost for specific categories')
add_bullet('Stage 2: Holiday/PCard Prioritization — Topic-specific result sorting')
add_bullet('Stage 3: Region Filtering — Multi-method country detection (metadata → filename → text → HR system names)')

# 6.6
doc.add_heading('6.6 Generation Agent', level=2)
add_para('File: agents/generation_agent.py', italic=True, color=GRAY)
add_para('Purpose: Final answer generation using Gemini with sophisticated prompt engineering.')
add_para('12 Critical Rules in Master Prompt:', bold=True)
add_bullet('Answer primary question FIRST (binary questions answered immediately)')
add_bullet('User region + entity prioritization')
add_bullet('OPCO labeling for each policy (HDS, GlobalLogic, HD)')
add_bullet('HV exclusion (only when explicitly queried)')
add_bullet('No greetings in answers — go straight to content')
add_bullet('Specific contacts over generic "contact HR"')
add_bullet('Honest gaps with fallback (AskNow + HRBP redirect)')
add_bullet('Professional tone')
add_bullet('STRICT country filtering (never cross-list multi-country policies)')
add_bullet('Answer in English (translation handled separately)')
add_bullet('Answer PRECISION (match the specific question)')
add_bullet('Source quality (prefer specific/recent documents)')
add_para('Modes: Concise (ultra-brief) vs Detailed (thorough with definitions, eligibility, rules, dates, processes, exceptions).')

# 6.7
doc.add_heading('6.7 Translation Agent', level=2)
add_para('File: agents/translation_agent.py', italic=True, color=GRAY)
add_bullet('LLM-based translation using Gemini with Markdown preservation')
add_bullet('Translation-only shortcut: "translate to Tamil" retranslates previous answer without re-running pipeline')
add_bullet('Fallback: If LLM fails, falls back to Google Cloud Translation API')

# 6.8
doc.add_heading('6.8 Post-Validation Agent', level=2)
add_para('File: agents/post_validation_agent.py', italic=True, color=GRAY)
add_para('Multi-stage validation pipeline:', bold=True)
add_bullet('Detect chitchat/greeting responses → no sources needed')
add_bullet('Detect "no information" fallback phrases')
add_bullet('Detect short redirect-only answers')
add_bullet('LLM source attribution: Ask Gemini which documents contributed to the answer')
add_bullet('Country mismatch filtering: Remove sources about wrong countries')
add_bullet('Fallback source extraction (score-based) if LLM attribution fails')

# 6.9
doc.add_heading('6.9 Supporting Agents', level=2)
add_table(
    ['Agent', 'Purpose'],
    [
        ['Caching Agent', 'Manages conversation history in Redis (save/retrieve/clear turns)'],
        ['Observability Agent', 'Structured logging: execution timing, user context boxes, agent step logs'],
        ['Feedback Agent', 'Stores user satisfaction ratings (helpful/unhelpful + comment) in Firestore'],
        ['Embedding Agent', 'Wrapper around Vertex AI text-embedding-004 for vectorization'],
    ],
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 7: Tools Layer
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Tools Layer', level=1)

doc.add_heading('7.1 BigQuery Tools (bq_tools.py)', level=2)
add_para('Critical for: User identity, roles, org hierarchy.', bold=True)
add_table(
    ['Function', 'Purpose'],
    [
        ['get_user_profile()', 'Consolidated call — roles + details + reports countries (2 parallel BQ queries, cached 24h)'],
        ['get_user_roles()', 'Extracts role flags from profile'],
        ['get_reportees_for_user()', 'Direct/2-level reports for employee lookup'],
        ['find_employee_in_reportees()', 'Search by name or LDAP ID'],
        ['search_employee_globally()', 'Search any employee (super admin only)'],
    ],
)
add_para('Key BQ Tables (in hd-onedata-prod):', bold=True)
add_bullet('hr_worker_snapshot_ds_hyperion_feed_hds_erp_vw — Manager relationships, country/region')
add_bullet('hr_worker_lifecycle_assignment_gblc_erp_vw — Job title, level, department, HR/Finance flags')
add_para('Role Detection Logic:', bold=True)
add_bullet('VP: Job title contains "vice"/"vp" OR job level in E2-E5 range')
add_bullet('HR: Oracle Department ID ∈ {252, 994}')
add_bullet('Finance: Oracle Department ID ∈ {312, 994}')

doc.add_heading('7.2 Embedding Tools (embedding_tools.py)', level=2)
add_bullet('Uses Vertex AI text-embedding-004 model')
add_bullet('Batches texts (100 per batch) with 3x retry and exponential backoff')
add_bullet('Task type: RETRIEVAL_DOCUMENT')

doc.add_heading('7.3 GCS Tools (gcs_tools.py)', level=2)
add_bullet('Upload/download/list blobs')
add_bullet('Signed URL generation: Smart path resolution for ServiceNow KB articles vs regular documents')
add_bullet('Cloud Run compatibility: IAM-based signing (signBlob API) instead of key-based')

doc.add_heading('7.4 OPCO Tools (opco_tools.py)', level=2)
add_para('Pure utility — no external dependencies.')
add_bullet('get_user_opco() — identify operating company from email domain')
add_bullet('is_hv_query() / is_hv_source() — detect Hitachi Vantara content')
add_bullet('is_holiday_query() — detect holiday calendar queries (with bereavement exclusions)')
add_bullet('is_p_card_query() — detect P-Card queries')
add_bullet('get_opco_entity() — classify document source (HDS, GL, HD, Global)')

doc.add_heading('7.5 Cache Tools (cache_tools.py)', level=2)
add_bullet('Redis wrapper with connection pooling, thread-safe operations')
add_bullet('Key-value, list, hash operations')
add_bullet('Semantic similarity cache: Stores (query, embedding, answer) triples')
add_bullet('Finds similar cached queries via cosine similarity (NumPy-accelerated)')

doc.add_heading('7.6 KB Renderer (kb_renderer.py)', level=2)
add_bullet('Converts ServiceNow KB extracted markdown into styled HTML with Hitachi Digital branding')
add_bullet('CSS stripping from extraction artifacts')
add_bullet('Table formatting (pipe-delimited → HTML tables)')
add_bullet('Deduplication of page-break fragments')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 8: Data Architecture
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Data Architecture & Index Groups', level=1)

doc.add_heading('Index Group Registry', level=2)
add_para('Each index group maps to a Vertex AI Matching Engine index + a Firestore database/collection:')
add_table(
    ['Index Group', 'Matching Engine Index', 'Firestore DB', 'Firestore Collection', 'Content'],
    [
        ['main', 'hd_skye2_0_*', 'hd-skye-db', 'hd-skye-2-0-chunks', 'Core HR policy docs (PDFs, DOCX)'],
        ['servicenow_kb', 'hd_skye_agents_servicenow', 'hd-skye-db-servicenow', 'hd-skye-chunks-servicenow', 'ServiceNow KB articles'],
        ['apac_payroll', 'apac_payroll_deployed', 'hd-skye-db-servicenow', 'apac-payroll-chunks', 'APAC payroll docs'],
        ['pcard', 'p_card_*', 'hd-skye-db-servicenow', 'p-card_policy', 'P-Card policy + table'],
        ['bulk_exp', 'bulk_exp_*', 'hd-skye-db-servicenow', 'bulk-expense', 'Bulk Expense policy docs'],
    ],
)

doc.add_heading('Firestore Document Structure (per chunk)', level=2)
add_code_block(
    '{\n'
    '  "id": "chunk_uuid",\n'
    '  "text": "The leave policy for India states...",\n'
    '  "source": "hd-skye-2.0/Documents/India_Leave_Policy.pdf",\n'
    '  "chunk_title": "Section 3.1",\n'
    '  "section_title": "Annual Leave Entitlement",\n'
    '  "country": "India",\n'
    '  "region": "APAC",\n'
    '  "category": "",\n'
    '  "document_type": "policy",\n'
    '  "language": "en",\n'
    '  "is_table": false,\n'
    '  "servicenow_url": "",\n'
    '  "servicenow_number": ""\n'
    '}'
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 9: ServiceNow Integration
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('9. ServiceNow Integration', level=1)

doc.add_heading('How KB Articles are Ingested', level=2)
add_bullet('KB articles are extracted from ServiceNow and stored as markdown files in GCS under servicenow_kb_extraction/ prefix', bold_prefix='Step 1 — Extraction: ')
add_bullet('The ingestion pipeline processes them like regular documents (parse → chunk → embed → store)', bold_prefix='Step 2 — Ingestion: ')
add_bullet('ServiceNow KB articles go into their own Matching Engine index and Firestore DB', bold_prefix='Step 3 — Separate Index: ')
add_bullet('scripts/sn_kb_crossref.json maps KB numbers to metadata', bold_prefix='Step 4 — Cross-reference: ')

doc.add_heading('How KB Articles are Served', level=2)
add_para('When a source is a ServiceNow KB article (pattern: ServiceNow_KB_KB*.html):')
add_bullet('Post-Validation builds the ServiceNow portal URL: https://hitachivantara.service-now.com/asknow?id=kb_article_view&sys_kb_id=...')
add_bullet('KB Renderer (/documents/ServiceNow_KB_*) fetches extracted markdown from GCS and renders styled HTML with Hitachi branding')
add_bullet('Source links point to either the ServiceNow portal URL or the rendered HTML page')

doc.add_heading('Category Boosting', level=2)
add_bullet('P-Card queries → boost ServiceNow KB articles with category="P Card"')
add_bullet('Bulk Expense queries → boost ServiceNow KB articles with category="Bulk Expense"')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 10: Ingestion Pipeline
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Document Ingestion Pipeline', level=1)

add_para('The ingestion pipeline (agents/ingestion_agent.py) handles end-to-end knowledge base population:')
add_bullet('PARSE — PDF: Document AI OCR (preferred) or pypdf fallback. DOCX: python-docx. Large PDFs: split into 15-page chunks for OCR.', bold_prefix='Step 1: ')
add_bullet('DETECT LANGUAGE — using langdetect library', bold_prefix='Step 2: ')
add_bullet('CHUNK — 1000 character chunks, 150 character overlap. Respects paragraph boundaries.', bold_prefix='Step 3: ')
add_bullet('EMBED — Vertex AI text-embedding-004. Batched (100/call), 3x retry with exponential backoff.', bold_prefix='Step 4: ')
add_bullet('STORE IN FIRESTORE — Chunk text + metadata + embedding vector per index group.', bold_prefix='Step 5: ')
add_bullet('UPSERT TO MATCHING ENGINE — Vector datapoints added to production index.', bold_prefix='Step 6: ')
add_para('Batch Ingestion: ingest_from_gcs() scans a GCS prefix, downloads each file, and runs the full pipeline.', italic=True)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 11: Caching Strategy
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Caching Strategy', level=1)

add_para('SKYE uses Redis for multiple caching layers:')
add_table(
    ['Layer', 'Key Pattern', 'TTL', 'Purpose'],
    [
        ['Conversation History', 'history:{session_id}', '24h', 'Multi-turn conversation memory'],
        ['Session Context', 'session:{session_id}:latest', '24h', 'User profile, roles, last Q&A'],
        ['Exact Answer Cache', 'answer_cache:{variant}:{query}:{region}', '3h', 'Exact-match response cache'],
        ['Semantic Similarity', 'sem_cache:{region}:{role_key}:*', '3h', 'Embedding-based similar query cache'],
        ['Search Cache', 'search:*', '1h', 'Vector search result cache'],
        ['User Profile Cache', 'user_profile_v1:{email}', '24h', 'BQ profile (roles, details)'],
        ['Reportees Cache', 'reportees_list_v3:{email}', '1h', 'Manager\'s direct reports'],
        ['Country Resolution', 'resolved_country_v3:{location}', '30d', 'Location → country mapping'],
    ],
)

doc.add_heading('Semantic Similarity Cache', level=2)
add_bullet('Generate embedding for incoming query')
add_bullet('Compare against cached query embeddings using cosine similarity (NumPy)')
add_bullet('If similarity ≥ 0.85 → return cached answer')
add_bullet('Bucketed by (region, role_key) to prevent cross-role leaks')
add_bullet('Intent validation: reject if cached intent ≠ query intent')

doc.add_heading('What\'s NOT Cached', level=2)
add_bullet('Greeting/follow-up/employee-lookup responses')
add_bullet('Denial/error/fallback messages ("I don\'t have information...")')
add_bullet('P-Card content in non-P-Card cache buckets')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 12: Access Control
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Access Control & Role-Based Architecture', level=1)

doc.add_heading('User Identity Flow', level=2)
add_para('When a request arrives with teams_metadata.email:')
add_bullet('get_user_profile(email) runs 2 parallel BigQuery queries (cached 24h)', bold_prefix='Step 1: ')
add_bullet('Derive roles: is_manager, is_vp, is_executive, is_hr, is_finance, is_super_admin', bold_prefix='Step 2: ')
add_bullet('Determine data_scope: "regional" (limited) or "global" (HR/GPS, overrides region check)', bold_prefix='Step 3: ')
add_bullet('Role-based query rewriting: augment search with role-specific terms', bold_prefix='Step 4: ')

doc.add_heading('Region Filtering Modes', level=2)
add_para('Controlled by REGION_FILTER_MODE env var:')
add_bullet('"all" (default) — All users get region-filtered results')
add_bullet('"managers_up" — Only Manager+ roles get region filtering')
add_bullet('"vp_up" — Only VP+ roles')
add_bullet('"none" — Disable for everyone')

doc.add_heading('Smart Region Heuristics', level=2)
add_bullet('Personal queries ("my leave", "am I entitled") → always use home location')
add_bullet('Bereavement queries → use home location, not mentioned location (event site ≠ policy country)')
add_bullet('Holiday queries → use home location for global-scope users')
add_bullet('Explicit global signals ("all countries", "compare") → keep global scope')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 13: Infrastructure
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Infrastructure & Deployment', level=1)

doc.add_heading('Docker Build (Multi-stage)', level=2)
add_bullet('Stage 1: Build React frontend with Bun → produces /dist with static assets')
add_bullet('Stage 2: Python 3.11-slim production image → pip install, copy frontend + backend, expose port 8000, run uvicorn')

doc.add_heading('Cloud Run Configuration', level=2)
add_table(
    ['Setting', 'Value'],
    [
        ['CPU', '2'],
        ['Memory', '2 GB'],
        ['Min instances', '0'],
        ['Max instances', '10'],
        ['Timeout', '300s'],
        ['Concurrency', '100 per instance'],
        ['CPU throttling', 'Disabled'],
        ['Session affinity', 'Enabled'],
        ['Authentication', 'Required'],
    ],
)

doc.add_heading('GCP Project Structure', level=2)
add_table(
    ['GCP Project', 'Purpose'],
    [
        ['hd-procurement-poc-gemini', 'App hosting (Cloud Run, Vertex AI, Firestore, GCS)'],
        ['hd-onedata-prod', 'User data (BigQuery HR worker snapshots)'],
    ],
)

doc.add_heading('Service Account', level=2)
add_para('hitachi-fin-service-account@hd-procurement-poc-gemini.iam.gserviceaccount.com')
add_para('Needs access to: Vertex AI, Firestore, GCS, Document AI, Cloud Translation, and cross-project BigQuery read on hd-onedata-prod.')

# ═════════════════════════════════════════════════════════════════════════
# SECTION 14: API Endpoints
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('14. API Endpoints', level=1)

doc.add_heading('Chat Endpoints', level=2)
add_table(
    ['Method', 'Path', 'Description', 'Variant'],
    [
        ['POST', '/chat', 'Main HR agent', 'main (or any via variant field)'],
        ['POST', '/pcard/chat', 'P-Card dedicated pipeline', 'pcard'],
        ['POST', '/bulk-expense/chat', 'Bulk Expense agent', 'bulk_expense'],
        ['POST', '/payroll/chat', 'Payroll agent', 'payroll'],
    ],
)

doc.add_heading('Request Body (QueryRequest)', level=2)
add_code_block(
    '{\n'
    '  "question": "What is the leave policy in India?",\n'
    '  "session_id": "user-123-session-456",\n'
    '  "teams_metadata": {"email": "john.doe@hitachidigital.com"},\n'
    '  "data_scope": "regional",\n'
    '  "variant": "main"\n'
    '}'
)

doc.add_heading('Response Body', level=2)
add_code_block(
    '{\n'
    '  "answer": "In India, employees are entitled to...",\n'
    '  "sources": ["India_Leave_Policy.pdf", "ServiceNow_KB_KB001234.html"],\n'
    '  "source_links": { ... },\n'
    '  "suggested_questions": ["What about sick leave?", ...],\n'
    '  "show_feedback_prompt": true,\n'
    '  "response_time_seconds": 4.23,\n'
    '  "variant": "main"\n'
    '}'
)

doc.add_heading('Utility Endpoints', level=2)
add_table(
    ['Method', 'Path', 'Description'],
    [
        ['GET', '/healthz', 'Health check'],
        ['POST', '/new-chat', 'Clear session history'],
        ['POST', '/feedback', 'Submit user feedback'],
        ['GET', '/documents/{filename}', 'Serve document (signed URL or rendered KB)'],
        ['GET', '/agents', 'List all registered agents'],
        ['GET', '/variants', 'List all agent variants + config'],
        ['GET', '/cache/status', 'Redis cache overview'],
        ['GET', '/cache/session/{id}', 'Session cache detail'],
    ],
)

# ═════════════════════════════════════════════════════════════════════════
# SECTION 15: Configuration Reference
# ═════════════════════════════════════════════════════════════════════════
doc.add_heading('15. Key Configuration Reference', level=1)

doc.add_heading('Critical Environment Variables', level=2)
add_table(
    ['Variable', 'Purpose', 'Example'],
    [
        ['PROJECT_ID', 'GCP project', 'hd-procurement-poc-gemini'],
        ['REGION', 'GCP region', 'us-central1'],
        ['LLM_MODEL', 'Gemini model', 'gemini-2.0-flash'],
        ['EMBEDDING_MODEL', 'Embedding model', 'text-embedding-004'],
        ['REDIS_HOST', 'Redis endpoint', '10.180.68.37'],
        ['SKYE_ENV', 'Environment selector', 'prod / poc'],
        ['ENABLED_INDEX_GROUPS', 'Active indexes', 'servicenow_kb,main,apac_payroll'],
        ['REGION_FILTER_MODE', 'Region filter scope', 'all / managers_up / none'],
        ['DEFAULT_VARIANT', 'Default agent variant', 'main'],
        ['INDEX_PRIORITY_BOOST', 'Score boost for priority indexes', '0.20'],
        ['SIMILARITY_THRESHOLD', 'Semantic cache threshold', '0.95'],
        ['LLM_THINKING_BUDGET', 'Gemini 2.5 thinking tokens', '0 (disabled)'],
    ],
)

doc.add_heading('Environment Files', level=2)
add_bullet('.env — Local development defaults')
add_bullet('.env.poc — POC project resources')
add_bullet('.env.prod — Production project resources')
add_bullet('env/prod.yaml — Cloud Run production env vars')
add_bullet('env/prod-poc.yaml — Cloud Run POC env vars')

# ═════════════════════════════════════════════════════════════════════════
# SECTION: File Structure
# ═════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Appendix: File Structure Summary', level=1)

add_code_block(
    'SKYE/\n'
    '├── main.py                     # FastAPI app + all HTTP endpoints\n'
    '├── config.py                   # Centralized config, variant registry, index groups\n'
    '├── Dockerfile                  # Multi-stage build (Bun frontend + Python backend)\n'
    '├── deploy.config.yaml          # Cloud Run deployment config\n'
    '├── requirements.txt            # Python dependencies\n'
    '│\n'
    '├── agents/\n'
    '│   ├── agent.py                # Agent registry (all 15 agents)\n'
    '│   ├── orchestrator.py         # Main pipeline orchestrator\n'
    '│   ├── pcard_orchestrator.py   # P-Card dedicated pipeline\n'
    '│   ├── query_understanding_agent.py\n'
    '│   ├── guardrails_agent.py\n'
    '│   ├── access_control_agent.py\n'
    '│   ├── retrieval_agent.py\n'
    '│   ├── reranking_agent.py\n'
    '│   ├── generation_agent.py\n'
    '│   ├── translation_agent.py\n'
    '│   ├── post_validation_agent.py\n'
    '│   ├── caching_agent.py\n'
    '│   ├── observability_agent.py\n'
    '│   ├── feedback_agent.py\n'
    '│   ├── embedding_agent.py\n'
    '│   ├── ingestion_agent.py\n'
    '│   ├── parsing_chunking_agent.py\n'
    '│   └── indexing_agent.py\n'
    '│\n'
    '├── tools/\n'
    '│   ├── bq_tools.py             # BigQuery user/role/org queries\n'
    '│   ├── embedding_tools.py      # Vertex AI embedding generation\n'
    '│   ├── gcs_tools.py            # Cloud Storage operations\n'
    '│   ├── opco_tools.py           # OPCO classification + query detection\n'
    '│   ├── cache_tools.py          # Redis cache + semantic similarity\n'
    '│   └── kb_renderer.py          # ServiceNow KB → styled HTML\n'
    '│\n'
    '├── frontend/                   # React SPA (Vite + Bun)\n'
    '├── scripts/                    # UAT evaluation scripts\n'
    '├── docs/                       # Documentation\n'
    '└── env/                        # Per-environment Cloud Run configs'
)

# ─── Save ────────────────────────────────────────────────────────────────
output_path = r'c:\Users\SSLTP11340\Desktop\SKYE\docs\KT_ARCHITECTURE.docx'
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
