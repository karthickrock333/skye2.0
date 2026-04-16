"""
parsing_chunking_agent.py
==========================
Handles document parsing (PDF, DOCX) and intelligent text chunking
for the ingestion pipeline.
"""
import os
import io
from google.adk.agents import Agent
from config import PROJECT_ID, DOCUMENT_AI_OCR_ID, DOCUMENT_AI_FORM_PARSER_ID, DOCUMENT_AI_LOCATION
import logging

logger = logging.getLogger("HD_SKYE_AGENT")


def read_document(file_path: str) -> str:
    """Read content from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        if DOCUMENT_AI_OCR_ID:
            return _read_pdf_docai(file_path)
        return _read_pdf_pypdf(file_path)
    elif ext == ".docx":
        return _read_docx(file_path)
    return ""


def _read_pdf_pypdf(path: str) -> str:
    import pypdf
    text = ""
    try:
        for page in pypdf.PdfReader(path).pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    except Exception as e:
        logger.error(f"Error reading PDF {path}: {e}")
    return text


def _read_docx(path: str) -> str:
    import docx
    text = ""
    try:
        for para in docx.Document(path).paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error reading DOCX {path}: {e}")
    return text


def _split_pdf(path: str, max_pages: int = 15) -> list:
    import pypdf
    chunks = []
    try:
        reader = pypdf.PdfReader(path)
        for i in range(0, len(reader.pages), max_pages):
            writer = pypdf.PdfWriter()
            for pn in range(i, min(i + max_pages, len(reader.pages))):
                writer.add_page(reader.pages[pn])
            buf = io.BytesIO()
            writer.write(buf)
            chunks.append(buf.getvalue())
            buf.close()
    except Exception as e:
        logger.error(f"Error splitting PDF: {e}")
    return chunks


def _read_pdf_docai(path: str) -> str:
    from google.cloud import documentai
    try:
        client = documentai.DocumentProcessorServiceClient()
        ocr_name = client.processor_path(PROJECT_ID, DOCUMENT_AI_LOCATION, DOCUMENT_AI_OCR_ID)
        combined = ""
        for chunk in _split_pdf(path):
            raw_doc = documentai.RawDocument(content=chunk, mime_type="application/pdf")
            result = client.process_document(request=documentai.ProcessRequest(name=ocr_name, raw_document=raw_doc))
            combined += result.document.text + "\n"
        return combined
    except Exception as e:
        logger.error(f"Document AI failed: {e}. Falling back to pypdf.")
        return _read_pdf_pypdf(path)


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> list:
    """Chunk text respecting paragraph boundaries."""
    if not text:
        return []
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) <= chunk_size:
            current += para + "\n\n"
        else:
            if current:
                chunks.append(current.strip())
            if len(para) > chunk_size:
                start = 0
                while start < len(para):
                    chunks.append(para[start : start + chunk_size])
                    start += chunk_size - overlap
                current = ""
            else:
                current = para + "\n\n"
    if current:
        chunks.append(current.strip())
    return chunks


def detect_language(text: str) -> str:
    if not text.strip():
        return "unknown"
    try:
        from langdetect import detect
        return detect(text)
    except Exception:
        return "unknown"


parsing_chunking_agent = Agent(
    name="parsing_chunking_agent",
    model="gemini-2.0-flash",
    description="Parses documents (PDF/DOCX) and chunks text for embedding.",
    instruction="Use read_document and chunk_text tools to process files for ingestion.",
    tools=[read_document, chunk_text, detect_language],
)
